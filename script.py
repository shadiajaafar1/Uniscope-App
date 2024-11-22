import streamlit as st
from langchain_openai import ChatOpenAI
from tavily import TavilyClient
from langchain_community.document_loaders import PyPDFLoader
import os
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
import json
from langchain.prompts import ChatPromptTemplate
from streamlit_option_menu import option_menu
from langchain.chains import create_retrieval_chain
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import UnstructuredExcelLoader
from langchain_community.document_loaders import UnstructuredCSVLoader
from dotenv import load_dotenv
from docx import Document as DocxDocument
from io import BytesIO

path = "documents/"
doc_list = os.listdir(path)

class Document:
    """
    Clase que representa un documento.
    
    Atributos:
    - page_content (str): Contenido de la página del documento.
    - metadata (dict): Metadatos del documento que incluyen el título y la URL.
    """

    def __init__(self, page_content, title, url):
        self.page_content = page_content
        self.metadata = {"title": title, "url": url}

# Función para cargar documentos
def load_documents(file_paths):
    documents = []
    
    for file_path in file_paths:
        try:
            # Detectar el tipo de archivo según la extensión
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == ".pdf":
                loader = PyPDFLoader(file_path)
                documents.append(loader.load())
                print("PDF Cargado Correctamente.")
            elif ext == ".csv":
                loader = UnstructuredCSVLoader(file_path=file_path)
                documents.append(loader.load())
                print("CSV cargado correctamente.")
            #elif ext in [".xls", ".xlsx"]:
              #  loader = UnstructuredExcelLoader(file_path=file_path, mode="elements")
              #  print(loader.load())
               # documents.extend(loader.load())  # `.load()` devuelve una lista de documentos
                #print("Excel cargado correctamente.")
            elif ext == ".txt":
                loader = TextLoader(file_path=file_path)
                documents.append(loader.load())  # `.load()` devuelve una lista de documentos
                print("TXT cargado correctamente.")
            else:
                print(f"Tipo de archivo no soportado: {file_path}")
                continue
            
            print(f"Archivo cargado: {file_path}")
        except Exception as e:
            print(f"Error al cargar el archivo {file_path}: {e}")
    
    return documents


def generate_search_queries(user_query):
            prompt_template = ChatPromptTemplate.from_messages([
                ("system", "Given a user query, generate 5 relevant search queries to retrieve data for analysis."),
                ("human", user_query)
            ])
            messages = prompt_template.format_messages(input=user_query)
            response = llm.invoke(messages)

            search_queries = response.content.strip().split("\n")
            with open("search_queries.json", "w") as f:
                json.dump({"queries": search_queries}, f)
            
            return search_queries

def web_search(queries):
            documents = []
            for query in queries:
                response = tavily_client.search(query)
                if "results" in response:
                    for result in response["results"]:
                        if "content" in result and result["content"]:
                            documents.append(Document(result["content"], result["title"], result["url"]))
            return documents

def create_word_document(text):
    doc = DocxDocument()
    doc.add_heading("Reporte Generado", level=1)
    doc.add_paragraph(text)
    return doc

def generate_and_download_word(answer):
    doc = create_word_document(answer)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Configurar el ambiente para OpenAI y Tavily----------------------------
load_dotenv()

os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
tavily_api_key = os.getenv("TAVILY_API_KEY")
tavily_client = TavilyClient(api_key=tavily_api_key)

llm = ChatOpenAI(model="gpt-4")

#Cargar documentos--------------------------------------------------------
file_paths = [os.path.join(path, doc) for doc in doc_list]
loaded_docs = load_documents(file_paths)
print(f"Documentos cargados: {len(loaded_docs)}")

#APP-----------------------------------------------------------------------
with st.sidebar:
    selected = option_menu(
        "Navigation",
        ["App", "Dashboard"],
        icons=["app-indicator", "bar-chart"],
        menu_icon="cast",
        default_index=0,
    )


if selected == "App":
    
    st.title("Ask UniScope")
    st.write("Expanding Prespectives on Graduate Education")

    if "uploaded_files_list" not in st.session_state:
        st.session_state.uploaded_files_list = []

    # Input de usuario
    prompt = st.text_area("Prompt", placeholder="Escribe tu prompt aquí...")
    uploaded_file = st.file_uploader("Opcional: Sube tu archivo aquí", type=["csv", "xlsx", "pdf", "txt"])
    
    if st.button("Enviar"):
        
        if uploaded_file is not None:
            file_name = uploaded_file.name
            file_type = file_name.split('.')[-1] if '.' in file_name else None

            file_path = os.path.join(path, file_name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            st.session_state.uploaded_files_list.append(file_name)
            st.write(f"Archivo '{file_name}' cargado correctamente.")
            
        else:
            st.write("No se ha subido ningún archivo.")
        

    if prompt:
        input = prompt
        
        #Generar 5 preguntas relacionadas con el input del usuario
        queries = generate_search_queries(input)
        
        #Hacer búsquedas web
        web_results = web_search(queries)   
        
        context = ''
        for data in queries:
            context = context + data + ", "

        # Crear el sistema de recuperación (RAG) si hay documentos
        if len(loaded_docs) > 0:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            splits = []
            for doc in loaded_docs:
                splits = splits + text_splitter.split_documents(doc)
            
            vectorstore = InMemoryVectorStore.from_documents(
                documents=splits, embedding=OpenAIEmbeddings()
            )
                    
            retriever = vectorstore.as_retriever()
        else:
            st.write("No documents to process.")


        # Crear y ejecutar cadena RAG
        system_prompt = (
            "You are a data scientist tasked with analyzing global trends in postgraduate education. "
            "You extract three key insights and question-answering tasks. Use the following pieces of retrieved context to answer "
            "the question. If you don't know the answer, say that you don't know. Use three sentences maximum and keep the "
            "answer concise."
            "\n\n"
            "{context}"
        )

        qa_prompt_template = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}")
        ])
        
        question_answer_chain = create_stuff_documents_chain(llm, qa_prompt_template)
        rag_chain = create_retrieval_chain(retriever, question_answer_chain)
        
        results = rag_chain.invoke({"input": f"{input}"})
        rag_result = results["context"][0].page_content
        
        results_combined = rag_result + context
        
        # Generar prompt adecuado para el tipo de pregunta del usuario -----------------------------------------
        prompt_generator = (
            "You will receive a prompt from a user, and your task is to reformulate it to maximize the clarity, relevance, "
            "and usefulness of the response. Follow these steps to create a new optimized prompt: \n"
            "1. Identify the Type of Question: Determine if the user’s prompt is a: \n"
            "- Definition or Explanation: The user is asking for clarification on a concept or idea.\n"
            "- Comparison: The user wants differences or similarities between concepts.\n"
            "- In-depth Analysis: The user seeks a deep analysis or advanced applications.\n"
            "2. Add Specific Context: If the concept is technical or specialized, add instructions to include additional context, "
            "background, or relevant assumptions to make the concept more understandable.\n"
            "3. Structure Instructions: Add format specifications that will optimize the response, such as:\n"
            "- 'Explain in simple terms' if the user might be new to the topic.\n"
            "- 'Provide examples and analogies' to enhance understanding.\n"
            "- 'Compare point-by-point' if the question is a comparison.\n"
            "4. Depth Level: Determine the level of depth that seems appropriate based on the user’s prompt (basic, intermediate, advanced) "
            "and include it in the new prompt.\n"
            "5. Always include a section for references that contains a list of the sources used in the response. \n"
            "6. Include statistical data where applicable to strengthen the response.\n"
            f"{input}\n"
        )

        # Generar el contenido optimizado del prompt
        prompt_generated = llm.invoke(prompt_generator).content
        
        #Generar referencias-------------------
        references = []
        # Referencias de los resultados web
        for doc in web_results:
            if "title" in doc.metadata and "url" in doc.metadata:
                reference = f"{doc.metadata['title']} ({doc.metadata['url']})"
                references.append(reference)

        # Referencias de los documentos cargados
        for file_path in file_paths:
            reference = f"Document: {os.path.basename(file_path)}, Path: {file_path}"
            references.append(reference)
            
        references = "\n".join(references)
        

        # Reporte final----------------------------------------------------------------------
        report_prompt = (
            "As a data analyst in the education sector, your task is to generate a report in Spanish using the following instructions: \n"
            f"{prompt_generated}\n"
            "Based on the following findings from the processed documents:\n"
            f"{results_combined}\n"
            "Ensure the report is clear and concise, using accessible and professional language. Include:\n"
            "- A reference section listing all the sources used in the response.\n"
            "- Any statistical data relevant to the findings.\n"
            "Use the following references as the basis for constructing your report, but ensure that only those explicitly utilized in the content are included in the references section:\n"
            f"{references}\n"
            "Incorporate these references into the report explicitly to ensure traceability and credibility."
        )

        # Generar el reporte final
        report = llm.invoke(report_prompt)

        # Mostrar la respuesta final
        answer = report.content
        st.write(answer)
        
        buffer = generate_and_download_word(answer)
        
        st.download_button(
            label="Descargar reporte",
            data=buffer,
            file_name="reporte.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        
        
elif selected == "Dashboard":
    power_bi_url = "https://app.powerbi.com/view?r=eyJrIjoiZTczOTU2MWQtNjY2Ni00MjVlLThhNzYtNGVjN2YxZGVmMzk2IiwidCI6ImJhYjBiNjc5LWJkNWYtNGZlOC1iNTE2LWM2YjhiMzE3Yzc4MiIsImMiOjR9&pageName=6df24e2b9d35065da98a"
    st.components.v1.html(
    f'''
    <div style="display: flex; justify-content: center; align-items: center; height: 100vh; width: 100vw;">
        <iframe width="900" height="500" src="{power_bi_url}" frameborder="0" allowFullScreen="true"></iframe>
    </div>
    ''',
    height=500, width=900
)